import streamlit as st
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import numpy as np

# ====================== CONFIG ======================
st.set_page_config(
    page_title="Saul Damon High School | Term Analysis",
    layout="wide",
    page_icon="📊"
)

# ====================== STYLING ======================
st.markdown("""
<style>
    .stApp {background-color:#F5F7FA; color:#1F2937; font-family:Arial;}
    .main-title {font-size:42px; text-align:center; font-weight:bold; color:#1E3A8A;}
    .sub-title {font-size:24px; text-align:center; margin-bottom:30px; color:#4B5563;}
    .section-header {font-size:24px; font-weight:700; color:#1E40AF; margin:25px 0 10px 0;}
    .insight-box {padding:15px; border-radius:8px; margin:10px 0;}
</style>
""", unsafe_allow_html=True)

# ====================== SIDEBAR ======================
with st.sidebar:
    st.header("📋 Settings")
    uploaded_file = st.file_uploader("Upload TERM TEMPLATE (Excel)", type=["xlsx"])
    term = st.selectbox("Select Term", ["Term 1", "Term 2", "Term 3", "Term 4"])
    chart_type = st.selectbox("Main Chart Type", ["Bar", "Stacked Bar", "Pie"])

# ====================== HEADER ======================
st.markdown('<p class="main-title">SAUL DAMON HIGH SCHOOL</p>', unsafe_allow_html=True)
st.markdown(f'<p class="sub-title">{term} Performance Analysis Dashboard</p>', unsafe_allow_html=True)

if not uploaded_file:
    st.info("👆 Please upload your Excel TERM TEMPLATE to begin analysis.")
    st.stop()

# ====================== DATA PROCESSING ======================
@st.cache_data
def process_data(file):
    df_raw = pd.read_excel(file, header=None)
    
    # Find grade section starts
    grade_starts = df_raw.index[df_raw[0].astype(str).str.contains("GRADE", na=False, case=False)].tolist()
    grades = ["Grade 9", "Grade 10", "Grade 11", "Grade 12"]
    
    grade_dfs = {}
    
    for i, grade in enumerate(grades):
        if i >= len(grade_starts):
            break
        start_idx = grade_starts[i] + 2  # Skip header row
        end_idx = grade_starts[i + 1] if i + 1 < len(grade_starts) else len(df_raw)
        
        gdf = df_raw.iloc[start_idx:end_idx].copy().reset_index(drop=True)
        gdf = gdf.dropna(how="all").reset_index(drop=True)
        
        if gdf.empty:
            continue
            
        # Define expected columns
        cols = ["SUBJECT", "AVERAGE MARK", "LEVEL 1", "LEVEL 2", "LEVEL 3",
                "LEVEL 4", "LEVEL 5", "LEVEL 6", "LEVEL 7", "TOTAL"]
        
        # Take only available columns
        gdf = gdf.iloc[:, :len(cols)]
        gdf.columns = cols[:gdf.shape[1]]
        
        # Clean subject names
        gdf["SUBJECT"] = gdf["SUBJECT"].astype(str).str.strip()
        gdf = gdf[gdf["SUBJECT"].str.len() > 2]  # Remove empty/invalid rows
        
        # Convert numeric columns
        numeric_cols = ["AVERAGE MARK"] + [f"LEVEL {i}" for i in range(1, 8)]
        for col in numeric_cols:
            if col in gdf.columns:
                gdf[col] = pd.to_numeric(gdf[col], errors="coerce").fillna(0)
        
        # Ensure TOTAL column
        level_cols = [f"LEVEL {i}" for i in range(1, 8) if f"LEVEL {i}" in gdf.columns]
        if "TOTAL" not in gdf.columns:
            gdf["TOTAL"] = gdf[level_cols].sum(axis=1)
        else:
            gdf["TOTAL"] = pd.to_numeric(gdf["TOTAL"], errors="coerce").fillna(0)
        
        grade_dfs[grade] = gdf
    
    return grade_dfs

grade_dfs = process_data(uploaded_file)

# ====================== CHART HELPERS ======================
def create_bar_chart(gdf, title):
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.barplot(x="AVERAGE MARK", y="SUBJECT", data=gdf.sort_values("AVERAGE MARK", ascending=False), 
                palette="Blues_d", ax=ax)
    ax.set_title(title, fontsize=14, pad=20)
    ax.set_xlabel("Average Mark (%)")
    ax.set_ylabel("")
    for container in ax.containers:
        ax.bar_label(container, fmt='%.1f')
    plt.tight_layout()
    return fig

def create_pie_chart(gdf, title, values_col="AVERAGE MARK"):
    fig, ax = plt.subplots(figsize=(8, 8))
    values = gdf[values_col]
    labels = gdf["SUBJECT"]
    ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, 
           colors=sns.color_palette("pastel", len(labels)))
    ax.set_title(title)
    ax.axis('equal')
    return fig

def save_fig_to_bytes(fig):
    img = BytesIO()
    fig.savefig(img, format='png', dpi=200, bbox_inches='tight')
    plt.close(fig)
    img.seek(0)
    return img

# ====================== PROFESSIONAL WORD REPORT ======================
def generate_professional_report(grade_dfs, term):
    doc = Document()
    
    # Header
    title = doc.add_heading(f"Saul Damon High School\n{term} Performance Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # School-wide summary
    doc.add_heading("Executive Summary", level=1)
    total_learners = 0
    overall_avg = 0
    total_subjects = 0
    
    for gdf in grade_dfs.values():
        if not gdf.empty:
            total_learners += int(gdf["TOTAL"].sum())
            overall_avg += gdf["AVERAGE MARK"].mean() * len(gdf)
            total_subjects += len(gdf)
    
    if total_subjects > 0:
        overall_avg = overall_avg / total_subjects
    
    p = doc.add_paragraph()
    p.add_run(f"Total Learners: ").bold = True
    p.add_run(f"{total_learners}")
    p = doc.add_paragraph()
    p.add_run(f"Overall Average Mark: ").bold = True
    p.add_run(f"{overall_avg:.1f}%")
    
    doc.add_paragraph("This report provides detailed insights into learner performance, subject-level analysis, "
                     "and actionable recommendations for academic improvement.")
    
    for grade, gdf in grade_dfs.items():
        if gdf.empty:
            continue
            
        doc.add_page_break()
        doc.add_heading(grade, level=1)
        
        # Grade metrics
        avg_mark = gdf["AVERAGE MARK"].mean()
        total_learners_grade = int(gdf["TOTAL"].sum())
        
        doc.add_paragraph(f"Average Mark: {avg_mark:.1f}%", style='List Bullet')
        doc.add_paragraph(f"Total Learners: {total_learners_grade}", style='List Bullet')
        doc.add_paragraph(f"Subjects Analyzed: {len(gdf)}", style='List Bullet')
        
        # Bar Chart
        fig = create_bar_chart(gdf, f"{grade} - Subject Average Marks")
        chart_img = save_fig_to_bytes(fig)
        doc.add_picture(chart_img, width=Inches(6.5))
        
        # Detailed Table
        doc.add_heading("Subject Performance Table", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Subject"
        hdr_cells[1].text = "Avg Mark"
        hdr_cells[2].text = "Total Learners"
        hdr_cells[3].text = "Pass Rate"
        
        for _, row in gdf.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["SUBJECT"])
            cells[1].text = f"{row['AVERAGE MARK']:.1f}%"
            cells[2].text = str(int(row["TOTAL"]))
            
            pass_rate = ((row["TOTAL"] - row.get("LEVEL 1", 0)) / row["TOTAL"] * 100) if row["TOTAL"] > 0 else 0
            cells[3].text = f"{pass_rate:.1f}%"
        
        # Pass/Fail Analysis
        gdf["FAILED"] = gdf.get("LEVEL 1", 0)
        gdf["PASSED"] = gdf["TOTAL"] - gdf["FAILED"]
        
        doc.add_heading("Pass / Fail Distribution", level=2)
        fig2, ax2 = plt.subplots(figsize=(10, 5))
        gdf.set_index("SUBJECT")[["FAILED", "PASSED"]].plot(kind="barh", stacked=True, ax=ax2, 
                                                            color=["#EF4444", "#10B981"])
        ax2.set_title(f"{grade} Pass/Fail by Subject")
        img2 = save_fig_to_bytes(fig2)
        doc.add_picture(img2, width=Inches(6.5))
    
    # Recommendations section
    doc.add_heading("Key Insights & Recommendations", level=1)
    doc.add_paragraph("• Focus additional support on subjects with averages below 50% or high failure rates (>25%).")
    doc.add_paragraph("• Implement targeted intervention programs for Level 1-3 learners.")
    doc.add_paragraph("• Celebrate and share best practices from high-performing subjects.")
    doc.add_paragraph("• Consider curriculum review where consistent underperformance is observed.")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ====================== GENERATE FULL REPORT ======================
full_report = generate_professional_report(grade_dfs, term)

st.download_button(
    "📄 Download Full Professional School Report (DOCX with Charts)",
    data=full_report,
    file_name=f"Saul_Damon_{term.replace(' ', '_')}_Full_Report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True
)

# ====================== DASHBOARD SECTIONS ======================
st.markdown("---")

for grade, gdf in grade_dfs.items():
    if gdf.empty:
        continue
        
    st.markdown(f'<p class="section-header">{grade}</p>', unsafe_allow_html=True)
    
    # Grade-specific report download
    grade_report = generate_professional_report({grade: gdf}, term)
    st.download_button(
        f"📥 Download {grade} Detailed Report",
        data=grade_report,
        file_name=f"{grade}_{term.replace(' ', '_')}_Report.docx",
        key=f"dl_{grade}"
    )
    
    # Key Metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Subjects", len(gdf))
    with col2:
        st.metric("Average Mark", f"{gdf['AVERAGE MARK'].mean():.1f}%")
    with col3:
        st.metric("Total Learners", int(gdf["TOTAL"].sum()))
    with col4:
        pass_rate = (1 - gdf.get("LEVEL 1", 0).sum() / gdf["TOTAL"].sum() * 100) if gdf["TOTAL"].sum() > 0 else 0
        st.metric("Overall Pass Rate", f"{pass_rate:.1f}%")
    
    # Main Chart
    st.subheader("📊 Subject Performance")
    if chart_type == "Bar":
        fig = create_bar_chart(gdf, f"{grade} Subject Averages")
        st.pyplot(fig)
    elif chart_type == "Pie":
        fig = create_pie_chart(gdf, f"{grade} Average Marks Distribution")
        st.pyplot(fig)
    else:  # Stacked Bar
        level_cols = [f"LEVEL {i}" for i in range(1, 8) if f"LEVEL {i}" in gdf.columns]
        fig, ax = plt.subplots(figsize=(10, 6))
        gdf.set_index("SUBJECT")[level_cols].plot(kind="barh", stacked=True, ax=ax, colormap="Blues")
        ax.set_title(f"{grade} - Level Distribution by Subject")
        st.pyplot(fig)
    
    st.dataframe(
        gdf[["SUBJECT", "AVERAGE MARK", "TOTAL"] + 
            [col for col in gdf.columns if "LEVEL" in col]].round(1),
        use_container_width=True
    )
    
    # Level Distribution Pies
    st.markdown('<p class="section-header">Level Distribution per Subject</p>', unsafe_allow_html=True)
    level_cols = [f"LEVEL {i}" for i in range(1, 8) if f"LEVEL {i}" in gdf.columns]
    cols = st.columns(3)
    for i, (_, row) in enumerate(gdf.iterrows()):
        with cols[i % 3]:
            levels = row[level_cols]
            if levels.sum() == 0:
                continue
            fig, ax = plt.subplots(figsize=(5, 5))
            ax.pie(levels, labels=level_cols, autopct='%1.1f%%', startangle=90, 
                   colors=sns.color_palette("Blues"))
            ax.set_title(row["SUBJECT"], fontsize=11)
            ax.axis('equal')
            st.pyplot(fig)
    
    # Pass/Fail
    st.markdown('<p class="section-header">Pass / Fail Analysis</p>', unsafe_allow_html=True)
    gdf = gdf.copy()
    gdf["FAILED"] = gdf.get("LEVEL 1", 0)
    gdf["PASSED"] = gdf["TOTAL"] - gdf["FAILED"]
    
    fig, ax = plt.subplots(figsize=(10, 5))
    gdf.set_index("SUBJECT")[["FAILED", "PASSED"]].plot(kind="barh", stacked=True, ax=ax,
                                                        color=["#EF4444", "#10B981"])
    ax.set_title(f"{grade} Pass/Fail by Subject")
    st.pyplot(fig)
    
    # Insights & Suggestions
    st.markdown('<p class="section-header">🔍 Insights & Recommendations</p>', unsafe_allow_html=True)
    
    for _, row in gdf.iterrows():
        if row["TOTAL"] == 0:
            continue
        fail_rate = (row["FAILED"] / row["TOTAL"]) * 100
        avg = row["AVERAGE MARK"]
        
        if fail_rate > 35:
            st.error(f"**{row['SUBJECT']}**: Critical - High failure rate ({fail_rate:.1f}%). "
                    f"Immediate intervention recommended (tutoring, curriculum review).")
        elif fail_rate > 20:
            st.warning(f"**{row['SUBJECT']}**: Concern - Failure rate {fail_rate:.1f}%. "
                      f"Consider extra classes or peer support.")
        elif avg < 55:
            st.warning(f"**{row['SUBJECT']}**: Below target average ({avg:.1f}%). Focus on foundational concepts.")
        else:
            st.success(f"**{row['SUBJECT']}**: Strong performance ({avg:.1f}%). Maintain momentum.")

st.caption("Saul Damon High School • Professional Academic Performance Dashboard • Generated with ❤️")
